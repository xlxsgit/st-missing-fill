import docx
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH

def strip_numbering(text):
    # Remove leading numbering like "4.2.1 " or "第4章 "
    text = re.sub(r'^(第[0-9一二三四五六七八九十]+章\s*)', '', text)
    text = re.sub(r'^[0-9]+(\.[0-9]+)*\s+', '', text)
    return text

def add_paragraph(doc, text, style=None, align=None):
    p = doc.add_paragraph(text, style=style)
    if align:
        p.alignment = align
    return p

def insert_modeling(src_docx_path, md_path, out_docx_path):
    # Load source docx (the AI draft)
    doc = docx.Document(src_docx_path)
    # Locate Chapter 4 heading (Heading 1 containing "第4章" or "第四章")
    start_idx = None
    end_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Heading 1' and ("第4章" in para.text or "第四章" in para.text):
            start_idx = i
        elif start_idx is not None and para.style.name == 'Heading 1' and ("第5章" in para.text or "第五章" in para.text or "结论" in para.text):
            end_idx = i
            break
    if start_idx is None:
        raise ValueError('Chapter 4 heading not found')
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    # Delete existing content between start_idx+1 and end_idx-1
    for _ in range(end_idx - start_idx - 1):
        p = doc.paragraphs[start_idx + 1]._element
        p.getparent().remove(p)
    # Insert new content after the heading
    insert_point = doc.paragraphs[start_idx]
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        if not line:
            i += 1
            continue
        if line.startswith('# '):
            add_paragraph(doc, strip_numbering(line[2:]), style='Heading 1')
        elif line.startswith('## '):
            add_paragraph(doc, strip_numbering(line[3:]), style='Heading 2')
        elif line.startswith('### '):
            add_paragraph(doc, strip_numbering(line[4:]), style='Heading 3')
        elif line.startswith('#### '):
            add_paragraph(doc, strip_numbering(line[5:]), style='Heading 4')
        elif line.startswith('$$'):
            # Block formula
            formula_lines = []
            if line == '$$':
                i += 1
                while i < len(lines) and lines[i].strip() != '$$':
                    formula_lines.append(lines[i].strip())
                    i += 1
            else:
                formula_lines.append(line.strip(' $'))
            formula_text = '\n'.join(formula_lines)
            p = add_paragraph(doc, formula_text, align=WD_ALIGN_PARAGRAPH.CENTER)
            for run in p.runs:
                run.font.name = 'Cambria Math'
                run.italic = True
        else:
            # Normal paragraph, replace image placeholders if any
            if '【图' in line:
                add_paragraph(doc, '[此处为图片]', align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                add_paragraph(doc, line)
        i += 1
    doc.save(out_docx_path)
    print(f'Created {out_docx_path}')

if __name__ == '__main__':
    insert_modeling('docs/1 硕士论文ai稿.docx', 'docs/tmp_ideas/chapter4_modeling.md', 'docs/1 硕士论文ai稿.docx')
