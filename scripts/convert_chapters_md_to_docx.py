#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path
from html import escape

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
UNORDERED_RE = re.compile(r"^[-*]\s+(.*)$")


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)  # 小四
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_formula(doc: Document, eq_text: str) -> None:
    p = doc.add_paragraph()
    xml = (
        f'<m:oMathPara {nsdecls("m")}>'
        f"<m:oMath><m:r><m:t>{escape(eq_text)}</m:t></m:r></m:oMath>"
        f"</m:oMathPara>"
    )
    p._element.append(parse_xml(xml))


def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(12)


def _add_code_block(doc: Document, code_lines: list[str]) -> None:
    for line in code_lines:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
            r.font.size = Pt(10.5)


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    _set_normal_style(doc)

    in_code = False
    in_math = False
    code_buf: list[str] = []
    math_buf: list[str] = []
    para_buf: list[str] = []

    def flush_para() -> None:
        if not para_buf:
            return
        text = " ".join(s.strip() for s in para_buf if s.strip()).strip()
        para_buf.clear()
        if text:
            _add_paragraph(doc, text)

    for raw in lines:
        line = raw.rstrip("\n")
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                _add_code_block(doc, code_buf)
                code_buf.clear()
                in_code = False
            else:
                code_buf.append(line)
            continue

        if in_math:
            if stripped == "$$":
                _add_formula(doc, " ".join(math_buf).strip())
                math_buf.clear()
                in_math = False
            else:
                math_buf.append(stripped)
            continue

        if stripped.startswith("```"):
            flush_para()
            in_code = True
            continue

        if stripped == "$$":
            flush_para()
            in_math = True
            continue

        if not stripped:
            flush_para()
            continue

        # Single-line block formula: $$ ... $$
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            flush_para()
            _add_formula(doc, stripped[2:-2].strip())
            continue

        # Heading
        mh = HEADING_RE.match(stripped)
        if mh:
            flush_para()
            level = min(len(mh.group(1)), 4)
            title = mh.group(2).strip()
            doc.add_heading(title, level=level)
            continue

        # Ordered list
        mo = ORDERED_RE.match(stripped)
        if mo:
            flush_para()
            p = doc.add_paragraph(mo.group(1).strip(), style="List Number")
            for r in p.runs:
                r.font.name = "Times New Roman"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                r.font.size = Pt(12)
            continue

        # Unordered list
        mu = UNORDERED_RE.match(stripped)
        if mu:
            flush_para()
            p = doc.add_paragraph(mu.group(1).strip(), style="List Bullet")
            for r in p.runs:
                r.font.name = "Times New Roman"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                r.font.size = Pt(12)
            continue

        para_buf.append(stripped)

    flush_para()
    if in_math and math_buf:
        _add_formula(doc, " ".join(math_buf).strip())
    if in_code and code_buf:
        _add_code_block(doc, code_buf)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    md_dir = root / "docs" / "body"
    out_dir = md_dir / "docxs"

    chapter_files = [md_dir / f"chapter{i}.md" for i in range(1, 6)]
    for md in chapter_files:
        if not md.exists():
            raise FileNotFoundError(f"Missing file: {md}")
        out = out_dir / f"{md.stem}.docx"
        markdown_to_docx(md, out)
        print(out)


if __name__ == "__main__":
    main()

