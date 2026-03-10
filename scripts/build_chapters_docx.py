#!/usr/bin/env python3
from __future__ import annotations

import re
from html import escape
from pathlib import Path

import latex2mathml.converter
import mathml2omml
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt
from lxml import etree


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
UNORDERED_RE = re.compile(r"^[-*]\s+(.*)$")
INLINE_MATH_RE = re.compile(r"(?<!\\)\$(.+?)(?<!\\)\$")
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS = {"m": MATH_NS}


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)  # 小四
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _normalize_latex(expr: str) -> str:
    text = expr.strip()
    if not text:
        return text

    # Flatten aligned-like environments to a single-line equation so converters
    # do not emit invalid XML for alignment markers (&).
    text = re.sub(r"\\begin\{aligned\*?\}", "", text)
    text = re.sub(r"\\end\{aligned\*?\}", "", text)
    text = text.replace("&", "")
    text = text.replace("\\\\", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _parse_omml(omml: str):
    try:
        root = parse_xml(omml)
        _tighten_omml_layout(root)
        return root
    except Exception:
        parser = etree.XMLParser(recover=True)
        root = etree.fromstring(omml.encode("utf-8"), parser=parser)
        if root is None:
            raise ValueError("Failed to parse OMML content.")
        _tighten_omml_layout(root)
        return root


def _tighten_omml_layout(root) -> None:
    def _accent_char_for_acc(ch: str) -> str:
        mapping = {
            "^": "̂",  # combining circumflex accent
            "~": "̃",  # combining tilde
            "¯": "̅",  # combining overline
            "̂": "̂",
            "̃": "̃",
            "̅": "̅",
        }
        return mapping.get(ch, ch)

    def _replace_with_acc(target, base_node, accent_char: str) -> None:
        acc = etree.Element(f"{{{MATH_NS}}}acc")
        acc_pr = etree.SubElement(acc, f"{{{MATH_NS}}}accPr")
        chr_el = etree.SubElement(acc_pr, f"{{{MATH_NS}}}chr")
        chr_el.set(qn("m:val"), _accent_char_for_acc(accent_char))
        e_el = etree.SubElement(acc, f"{{{MATH_NS}}}e")
        for child in list(base_node):
            e_el.append(child)
        parent = target.getparent()
        if parent is not None:
            parent.replace(target, acc)

    # Convert accent-like limUpp forms to acc, matching Word native accent rendering.
    for limupp in list(root.xpath(".//m:limUpp", namespaces=NS)):
        lim_chars = "".join(limupp.xpath(".//m:lim//m:t/text()", namespaces=NS)).strip()
        if lim_chars not in {"^", "~", "¯"}:
            continue
        base = limupp.find("m:e", namespaces=NS)
        if base is None:
            continue
        _replace_with_acc(limupp, base, lim_chars)

    # Convert accent groupChr to acc to tighten visual distance.
    for group_chr in list(root.xpath(".//m:groupChr", namespaces=NS)):
        ch = "".join(group_chr.xpath("./m:groupChrPr/m:chr/@m:val", namespaces=NS)).strip()
        if ch not in {"^", "~", "¯", "̂", "̃", "̅"}:
            continue
        base = group_chr.find("m:e", namespaces=NS)
        if base is None:
            continue
        _replace_with_acc(group_chr, base, ch)

    # Remove unnecessary box wrappers so accents are visually closer to the base symbol.
    changed = True
    while changed:
        changed = False
        for box in list(root.xpath(".//m:box", namespaces=NS)):
            parent = box.getparent()
            if parent is None:
                continue
            e_el = box.find("m:e", namespaces=NS)
            if e_el is None:
                continue
            idx = parent.index(box)
            for child in list(e_el):
                parent.insert(idx, child)
                idx += 1
            parent.remove(box)
            changed = True


def _latex_to_omml(latex_expr: str) -> str:
    mathml = latex2mathml.converter.convert(_normalize_latex(latex_expr))
    omml = mathml2omml.convert(mathml)
    if "<m:oMath" in omml and "xmlns:m=" not in omml:
        omml = omml.replace("<m:oMath>", f"<m:oMath {nsdecls('m')}>", 1)
    return omml


def _add_display_formula(doc: Document, latex_expr: str) -> None:
    p = doc.add_paragraph()
    expr = latex_expr.strip()
    if not expr:
        return
    try:
        omml = _latex_to_omml(expr)
        # Wrap display formula as oMathPara for better Word rendering.
        inner = omml.strip()
        if inner.startswith("<m:oMath "):
            wrapped = f"<m:oMathPara {nsdecls('m')}>{inner}</m:oMathPara>"
        elif inner.startswith("<m:oMath>"):
            inner = inner.replace("<m:oMath>", f"<m:oMath {nsdecls('m')}>", 1)
            wrapped = f"<m:oMathPara {nsdecls('m')}>{inner}</m:oMathPara>"
        else:
            wrapped = f"<m:oMathPara {nsdecls('m')}><m:oMath><m:r><m:t>{escape(expr)}</m:t></m:r></m:oMath></m:oMathPara>"
        p._element.append(_parse_omml(wrapped))
    except Exception:
        # Fallback: keep formula text when conversion fails.
        r = p.add_run(expr)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(12)


def _add_inline_mixed_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    pos = 0
    for m in INLINE_MATH_RE.finditer(text):
        if m.start() > pos:
            seg = text[pos : m.start()]
            r = p.add_run(seg)
            r.font.name = "Times New Roman"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(12)

        expr = m.group(1).strip()
        try:
            omml = _latex_to_omml(expr)
            p._element.append(_parse_omml(omml))
        except Exception:
            r = p.add_run(m.group(0))
            r.font.name = "Times New Roman"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(12)
        pos = m.end()

    if pos < len(text):
        seg = text[pos:]
        r = p.add_run(seg)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(12)


def _add_code_block(doc: Document, code_lines: list[str]) -> None:
    for line in code_lines:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
            r.font.size = Pt(10.5)


def _append_markdown(doc: Document, md_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()

    in_code = False
    in_math = False
    code_buf: list[str] = []
    math_buf: list[str] = []
    para_buf: list[str] = []

    def flush_para() -> None:
        if not para_buf:
            return
        text = " ".join(s.strip() for s in para_buf if s.strip()).strip()
        para_buf.clear()
        if text:
            _add_inline_mixed_paragraph(doc, text)

    for raw in lines:
        stripped = raw.strip()

        if in_code:
            if stripped.startswith("```"):
                _add_code_block(doc, code_buf)
                code_buf.clear()
                in_code = False
            else:
                code_buf.append(raw.rstrip("\n"))
            continue

        if in_math:
            if stripped == "$$":
                _add_display_formula(doc, " ".join(math_buf))
                math_buf.clear()
                in_math = False
            else:
                math_buf.append(stripped)
            continue

        if stripped.startswith("```"):
            flush_para()
            in_code = True
            continue

        if stripped == "$$":
            flush_para()
            in_math = True
            continue

        if not stripped:
            flush_para()
            continue

        # Single-line display math: $$ ... $$
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            flush_para()
            _add_display_formula(doc, stripped[2:-2])
            continue

        mh = HEADING_RE.match(stripped)
        if mh:
            flush_para()
            level = min(len(mh.group(1)), 4)
            title = mh.group(2).strip()
            doc.add_heading(title, level=level)
            continue

        mo = ORDERED_RE.match(stripped)
        if mo:
            flush_para()
            p = doc.add_paragraph(style="List Number")
            _add_inline_mixed_paragraph_to_existing(p, mo.group(1).strip())
            continue

        mu = UNORDERED_RE.match(stripped)
        if mu:
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_mixed_paragraph_to_existing(p, mu.group(1).strip())
            continue

        para_buf.append(stripped)

    flush_para()
    if in_math and math_buf:
        _add_display_formula(doc, " ".join(math_buf))
    if in_code and code_buf:
        _add_code_block(doc, code_buf)


def _add_inline_mixed_paragraph_to_existing(p, text: str) -> None:
    pos = 0
    for m in INLINE_MATH_RE.finditer(text):
        if m.start() > pos:
            seg = text[pos : m.start()]
            r = p.add_run(seg)
            r.font.name = "Times New Roman"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(12)
        expr = m.group(1).strip()
        try:
            omml = _latex_to_omml(expr)
            p._element.append(_parse_omml(omml))
        except Exception:
            r = p.add_run(m.group(0))
            r.font.name = "Times New Roman"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(12)
        pos = m.end()
    if pos < len(text):
        seg = text[pos:]
        r = p.add_run(seg)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(12)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    md_dir = root / "docs" / "body"
    out_path = md_dir / "docxs" / "chapters.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_normal_style(doc)

    chapter_files = [md_dir / f"chapter{i}.md" for i in range(1, 6)]
    for idx, md in enumerate(chapter_files):
        if not md.exists():
            raise FileNotFoundError(f"Missing file: {md}")
        _append_markdown(doc, md)
        if idx != len(chapter_files) - 1:
            doc.add_page_break()

    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
